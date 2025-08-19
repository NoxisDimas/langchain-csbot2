import os
import json
import pandas as pd
from typing import List, Dict, Any, Optional
from pathlib import Path
import aiofiles
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangChainDocument
import pypdf
from docx import Document as DocxDocument
import markdown
import uuid
from datetime import datetime

from app.services.database_service import DatabaseService

class DocumentService:
    def __init__(self):
        # Legacy DB service no longer required for RAG-only flow
        try:
            self.db_service = DatabaseService()
        except Exception:
            self.db_service = None
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
        self.upload_dir = Path("uploads")
        self.upload_dir.mkdir(exist_ok=True)
    
    def get_file_extension(self, filename: str) -> str:
        """Get file extension from filename"""
        return Path(filename).suffix.lower()
    
    def is_supported_file(self, filename: str) -> bool:
        """Check if file type is supported"""
        supported_extensions = {
            '.pdf', '.txt', '.md', '.csv', '.xlsx', '.xls', 
            '.docx', '.doc', '.json', '.xml', '.html'
        }
        return self.get_file_extension(filename) in supported_extensions
    
    async def save_uploaded_file(self, file_content: bytes, filename: str) -> str:
        """Save uploaded file to disk"""
        # sanitize filename
        safe_name = os.path.basename(filename).replace("..", "").replace("/", "_").replace("\\", "_")
        if not safe_name:
            safe_name = f"upload_{uuid.uuid4().hex}"
        file_path = self.upload_dir / safe_name
        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(file_content)
        return str(file_path)

    def process_file_to_chunks(self, file_content: bytes, filename: str, knowledge_base_name: str = "default") -> List[LangChainDocument]:
        """Process uploaded file and return LangChain chunks with metadata (RAG-only flow)."""
        # Save to disk (for audit/debug)
        file_path = None
        try:
            file_path = self.upload_dir / os.path.basename(filename).replace("..", "").replace("/", "_").replace("\\", "_")
            with open(file_path, 'wb') as f:
                f.write(file_content)
        except Exception:
            pass

        file_type = self.get_file_extension(filename)
        text_content = self.extract_text_from_file(str(file_path) if file_path else filename, file_type)
        if not text_content.strip():
            return []
        chunks = self.chunk_text(text_content, {
            "filename": filename,
            "knowledge_base": knowledge_base_name,
        })
        # ensure chunk_index in metadata
        for idx, c in enumerate(chunks):
            c.metadata = {**(c.metadata or {}), "chunk_index": idx}
        return chunks
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")
            return ""
    
    def extract_text_from_csv(self, file_path: str) -> str:
        """Extract text from CSV file"""
        try:
            df = pd.read_csv(file_path)
            return df.to_string()
        except Exception as e:
            print(f"Error extracting text from CSV: {e}")
            return ""
    
    def extract_text_from_excel(self, file_path: str) -> str:
        """Extract text from Excel file"""
        try:
            df = pd.read_excel(file_path)
            return df.to_string()
        except Exception as e:
            print(f"Error extracting text from Excel: {e}")
            return ""
    
    def extract_text_from_markdown(self, file_path: str) -> str:
        """Extract text from Markdown file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                md_content = file.read()
                # Convert markdown to plain text
                html = markdown.markdown(md_content)
                # Simple HTML to text conversion
                import re
                text = re.sub(r'<[^>]+>', '', html)
                return text
        except Exception as e:
            print(f"Error extracting text from Markdown: {e}")
            return ""
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            print(f"Error extracting text from TXT: {e}")
            return ""
    
    def extract_text_from_file(self, file_path: str, file_type: str) -> str:
        """Extract text from file based on file type"""
        file_type = file_type.lower()
        
        if file_type == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_type == '.docx':
            return self.extract_text_from_docx(file_path)
        elif file_type == '.csv':
            return self.extract_text_from_csv(file_path)
        elif file_type in ['.xlsx', '.xls']:
            return self.extract_text_from_excel(file_path)
        elif file_type == '.md':
            return self.extract_text_from_markdown(file_path)
        elif file_type == '.txt':
            return self.extract_text_from_txt(file_path)
        else:
            # Try as text file
            return self.extract_text_from_txt(file_path)
    
    def chunk_text(self, text: str, metadata: Dict[str, Any] = None) -> List[LangChainDocument]:
        """Split text into chunks"""
        if not text.strip():
            return []
        
        # Create LangChain documents
        docs = [LangChainDocument(page_content=text, metadata=metadata or {})]
        
        # Split into chunks
        chunks = self.text_splitter.split_documents(docs)
        return chunks
    
    async def process_file(self, file_content: bytes, filename: str, knowledge_base_name: str = "default") -> Optional[Document]:
        """Process uploaded file and store in database"""
        try:
            # Check if file is supported
            if not self.is_supported_file(filename):
                raise ValueError(f"Unsupported file type: {self.get_file_extension(filename)}")
            
            # Save file to disk
            file_path = await self.save_uploaded_file(file_content, filename)
            file_size = len(file_content)
            file_type = self.get_file_extension(filename)
            
            # Extract text from file
            text_content = self.extract_text_from_file(file_path, file_type)
            
            if not text_content.strip():
                raise ValueError("No text content extracted from file")
            
            # Create or get knowledge base (idempotent)
            kb = self.db_service.get_knowledge_base(knowledge_base_name)
            if not kb:
                kb = self.db_service.create_knowledge_base(knowledge_base_name)
                if not kb:
                    # If creation failed due to race or uniqueness, attempt to fetch again
                    kb = self.db_service.get_knowledge_base(knowledge_base_name)
                    if not kb:
                        raise ValueError("Failed to create knowledge base")

            # Ensure tables and column migrations are applied for this KB schema
            try:
                self.db_service.create_tables(kb.schema_name)
            except Exception:
                pass
            
            # Create document record
            with self.db_service.get_session() as session:
                document = Document(
                    filename=filename,
                    file_path=file_path,
                    file_type=file_type,
                    file_size=file_size,
                    content=text_content,
                    document_metadata=json.dumps({
                        "knowledge_base": knowledge_base_name,
                        "processed_at": datetime.utcnow().isoformat(),
                        "file_type": file_type
                    }),
                    is_processed=False
                )
                session.add(document)
                session.commit()
                session.refresh(document)
            
            # Create chunks
            chunks = self.chunk_text(text_content, {
                "document_id": str(document.id),
                "filename": filename,
                "knowledge_base": knowledge_base_name
            })
            
            # Store chunks (embeddings will be created later)
            with self.db_service.get_session() as session:
                for i, chunk in enumerate(chunks):
                    embedding = DocumentEmbedding(
                        document_id=document.id,
                        chunk_index=i,
                        chunk_text=chunk.page_content,
                        embedding_metadata=json.dumps(chunk.metadata)
                    )
                    session.add(embedding)
                
                # Mark document as processed
                document.is_processed = True
                session.commit()
            
            return document
            
        except Exception as e:
            print(f"Error processing file: {e}")
            return None
    
    def get_document_chunks(self, document_id: str) -> List[DocumentEmbedding]:
        """Get all chunks for a document"""
        with self.db_service.get_session() as session:
            return session.query(DocumentEmbedding).filter(
                DocumentEmbedding.document_id == document_id
            ).order_by(DocumentEmbedding.chunk_index).all()
    
    def search_documents(self, query: str, knowledge_base_name: str = None, limit: int = 10) -> List[Dict[str, Any]]:
        """Search documents by query (placeholder for vector search)"""
        # This will be implemented with vector search
        with self.db_service.get_session() as session:
            query_obj = session.query(Document)
            if knowledge_base_name:
                # Filter by knowledge base if specified
                pass
            
            documents = query_obj.limit(limit).all()
            return [
                {
                    "id": str(doc.id),
                    "filename": doc.filename,
                    "file_type": doc.file_type,
                    "created_at": doc.created_at.isoformat(),
                    "is_processed": doc.is_processed
                }
                for doc in documents
            ]